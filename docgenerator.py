import docx
import pandas as pd
from datetime import date, datetime
import calendar
from app import get_reports_month_rendered_hours, format_timedelta, get_month_info, get_reports_month_weekly, get_data,get_all_attendance_month,getMemberNamePosition,getFormatTime
from docx.shared import Pt  



def generatedoc(month,year, filter_member):
    # Initialise the Word document
    doc = docx.Document()
    month = int(month)
    year = int(year)
    month_totalhours = get_reports_month_rendered_hours(month, year, filter_member)
    month_totalhours = list(month_totalhours)
    
    members = []
    total_rendered = []
    remarks = []
    position = []
    for each in month_totalhours:
        members.append(each['name'])
        total_rendered.append(str(format_timedelta(each['total_rendered'])))
        remarks.append(each['remarks'])
        position.append(each['position'])
    
    
    
    data = {
        'Name' : members,
        'Position':position,
        'Total Rendered Hours' : total_rendered,
        'Remarks': remarks
    }

    df = pd.DataFrame(data)
    # print(df)
    word_month = calendar.month_name[int(month)]

    doc.add_heading(f'Lakandiwa {word_month} {year} DTR Report', 0)

    member_name = None
    if filter_member is not None and filter_member != 'all':
        member_name = get_data('member','id_number', filter_member)

    p = doc.add_paragraph(f"These are the list of Lakandiwa members and their rendered hours of duty for the month of {word_month} {year}." if filter_member == 'all' else f"These are the list of {member_name['firstname']}'s rendered hours of duty for the month of {word_month} {year}.")

    # Initialise the table
    t = doc.add_table(rows=1, cols=df.shape[1])

    # Add borders
    t.style = 'TableGrid'

    # Add the column headings
    for j in range(df.shape[1]):
        t.cell(0, j).text = df.columns[j]
    # Add the body of the data frame
    for i in range(df.shape[0]):
        row = t.add_row()
        for j in range(df.shape[1]):
            cell = df.iat[i, j]
            row.cells[j].text = str(cell)


    month_weekly_totalhours = get_reports_month_weekly(month, year, filter_member)
    month_weekly_totalhours = list(month_weekly_totalhours)
    month_info = get_month_info(month, year)
    doc.add_heading(f'Weekly Attendance', 0)
     
    pa = doc.add_paragraph(f"These are the list of Lakandiwa members based on their weekly rendered hours of duty for the month of {word_month} {year}." if filter_member == 'all' else f"These are the list of {member_name['firstname']}'s weekly rendered hours of duty for the month of {word_month} {year}.")
    for info in month_info:
        doc.add_paragraph(f'Week '+ str(info['week_number']) +' Report | '+str(info['week_start']) +' - '+ str(info['week_end']))
        weekly_members = []
        weekly_total_rendered = []
        weekly_remarks = []
        weekly_position = []
        for each in month_weekly_totalhours:
            if each['week_number'] == info['week_number']:
                weekly_members.append(each['member_name'])
                weekly_total_rendered.append(str(format_timedelta(each['time_total'])))
                weekly_remarks.append(each['remarks'])
                weekly_position.append(each['position'])
        weekly_data = {
            'Name' : weekly_members,
            'Position':weekly_position,
            'Week Total Rendered Hours' : weekly_total_rendered,
            'Remarks': weekly_remarks
        }

        
        wdf = pd.DataFrame(weekly_data)
    

       

        # Initialise the table
        ta = doc.add_table(rows=1, cols=df.shape[1])

        # Add borders
        ta.style = 'TableGrid'

        # Add the column headings
        for j in range(wdf.shape[1]):
            ta.cell(0, j).text = wdf.columns[j]
        # Add the body of the data frame
        for i in range(wdf.shape[0]):
            row = ta.add_row()
            for j in range(wdf.shape[1]):
                cell = wdf.iat[i, j]
                row.cells[j].text = str(cell)
    
    daily_attendance = None
    if filter_member is not None and filter_member != 'all':
        daily_attendance = get_all_attendance_month(month, year, filter_member)

    if daily_attendance:    
        member_time_in = []
        member_countersign_time_in = []
        member_signature_time_in = []
        member_time_out = []
        member_countersign_time_out = []
        member_signature_time_out = []
        member_work_done = []
        member_time_rendered = []

        for each in daily_attendance:
            member_time_in.append(getFormatTime(each['time_in']))
            member_countersign_time_in.append(getMemberNamePosition(each['countersign_time_in']))
            member_signature_time_in.append(getMemberNamePosition(each['signature_time_in']))
            member_time_out.append(getFormatTime(each['time_out']))
            member_countersign_time_out.append(getMemberNamePosition(each['countersign_time_out']))
            member_signature_time_out.append(getMemberNamePosition(each['signature_time_out']))
            member_work_done.append(each['work_done'])
            member_time_rendered.append(str(format_timedelta(each['time_rendered'])))

        daily_data = {
            'Time in': member_time_in,
            'Countersign Time in':member_countersign_time_in,
            'Signature Time in': member_signature_time_in,
            'Time out':member_time_out,
            'Countersign Time out':member_countersign_time_out,
            'Signature Time out': member_signature_time_out,
            'Work done':member_work_done,
            'Time rendered':member_time_rendered
        }

        doc.add_heading(f'Daily Attendance', 0)
        paa = doc.add_paragraph(f"These are the list of {member_name['firstname']}'s daily attendance the month of {word_month} {year}.")
        ddf = pd.DataFrame(daily_data)
    

       

        # Initialise the table
        tta = doc.add_table(rows=1, cols=ddf.shape[1])

        # Add borders
        tta.style = 'TableGrid'

        # Add the column headings
        for j in range(ddf.shape[1]):
            tta.cell(0, j).text = ddf.columns[j]
        # Add the body of the data frame
        for i in range(ddf.shape[0]):
            row = tta.add_row()
            for j in range(ddf.shape[1]):
                cell = ddf.iat[i, j]
                # row.cells[j].text = str(cell)
                cell_text = str(cell)
                cell_paragraph = row.cells[j].paragraphs[0]
                cell_paragraph.text = cell_text
                for run in cell_paragraph.runs:
                    font = run.font
                    font.size = Pt(8)
                
        
    # Save the Word doc
    filter_string = ''
    if member_name:
        filter_string = f'{member_name["lastname"].upper()}_{member_name["firstname"].upper()}_'
    doc.save(f'reports/{filter_string if filter_member != "all" else "" }{word_month}_{year}_DTR_report.docx')